from flask import Flask, session, render_template, request, flash, redirect, after_this_request
from flask_session import Session
from tempfile import mkdtemp
from werkzeug.utils import secure_filename
import ocrspace
import os
from docx import Document


UPLOAD_FOLDER = 'static/uploads'
FILES_DIR = 'static/texts'
FILENAME = ''


ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg'}
langs = [
    'Arabic',
    'Bulgarian',
    'Chinese_Simplified',
    'Chinese_Traditional',
    'Croatian',
    'Danish',
    'Dutch',
    'English',
    'Finnish',
    'French',
    'German',
    'Greek',
    'Hungarian',
    'Italian',
    'Japanese',
    'Korean',
    'Norwegian',
    'Polish',
    'Portuguese',
    'Russian',
    'Slovenian',
    'Spanish',
    'Swedish',
    'Turkish'
]

app = Flask( __name__)

# Ensure templates are auto-reloaded
app.config["TEMPLATES_AUTO_RELOAD"] = True

# link static files
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['FILES_DIR'] = FILES_DIR

# utilze
app.secret_key = 'super secret key'


# Configure session to use filesystem (instead of signed cookies)
app.config["SESSION_FILE_DIR"] = mkdtemp()
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"
Session(app)

# main or home page
@app.route("/")
def home(msg = None, type = 'warning'):
    
    #session['id'] = request.cookies.get('session')

    # check if others def rutrun alert
    if msg:
        msg = alert(msg, type)
    else:
        msg = False
    
    return render_template('home.html', home = 'active', langs = langs, alert = msg)


@app.route("/about")
def about(title = None):
    
    return render_template('about.html', about = 'active' )


"""
Editing text that is extract from file here and also chosse to save as Word doc or txt file
"""
@app.route("/editing", methods=["GET", "POST"])
def editing(title = None):
    
    if request.method == 'POST':

        # get file
        file = request.files['file']
        
        # check if file extention allowed
        if not(file and allowed_file(file.filename)):
            return home('PDF, PNG, JPG and JPEG files are allowed!')
        
        # secure file name and save
        filename = secure_filename(file.filename)
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

        # get file uploaded name
        full_filename = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        try:
            # connect to API
            api = eval(f"ocrspace.API('e36c320cbb88957', ocrspace.Language.{request.form.get('lang')})")
            
            # get text from server
            text = api.ocr_file(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        except Exception as e:
            return home(e)

        finally:
            # remove file after process
            os.remove(full_filename)
        
        
        

         
        
        # return result
        return render_template('editing.html', text=text)
    # return alert if file isnt upload
    return home('Should upload file first!')


# show final text and download file
@app.route("/finish", methods=["GET", "POST"])
def finish(title = None):
    # get userID session
    session['id'] = request.cookies.get('session')
    if request.method == 'POST':

        # check if user want to download as docs or txt
        if request.form.get('type') == 'doc':   # if doc
            doc = Document()                                                                # create new Doc
            doc.add_heading('Text extract from file', 0)                                    # add heading
            doc.add_paragraph(request.form.get('text'))                                     # add text
            doc.save(os.path.join(app.config['FILES_DIR'], session['id']+'.docx'))          # save file
            FILENAME = '.docx'                                                              # save extecntion
        else:       # if txt
            txt = open(os.path.join(app.config['FILES_DIR'], session['id']+'.txt'), 'w')    # create new txt
            txt.write(request.form.get('text'))                                             # add heading
            txt.close                                                                       # save file 
            FILENAME = '.txt'                                                               # save extecntion
        FILENAME = os.path.join(app.config['FILES_DIR'], session['id']+ FILENAME)           # get file
        return render_template('finish.html', file=FILENAME, text=request.form.get('text')) 

    # return alert if file isnt upload
    return home('Should upload file first!')

def alert(text, type='primary'):
    return {
        'msg': text,
        'type': type
    }

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

